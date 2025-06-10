import os
import json
import shutil
from datetime import datetime
import zipfile
import pandas as pd
from docx import Document
from docx.shared import Pt


def setup_data_files():
    """Cria os arquivos e diretórios necessários se não existirem"""
    os.makedirs('data', exist_ok=True)
    os.makedirs('assets/templates', exist_ok=True)

    # Cria config.json se não existir
    if not os.path.exists('data/config.json'):
        with open('data/config.json', 'w', encoding='utf-8') as f:
            json.dump({
                'theme': 'dark',
                'default_dir': os.path.expanduser('~/Documents/Projects'),
                'recent_projects': [],
                'pomodoro_duration': 25,
                'font_size': 10,
                'font_family': 'Segoe UI' if os.name == 'nt' else 'Helvetica'
            }, f, indent=4, ensure_ascii=False)

    # Cria historico.json se não existir
    if not os.path.exists('data/historico.json'):
        with open('data/historico.json', 'w', encoding='utf-8') as f:
            json.dump([], f, ensure_ascii=False)

    # Cria notes.json se não existir
    if not os.path.exists('data/notes.json'):
        with open('data/notes.json', 'w', encoding='utf-8') as f:
            json.dump([], f, ensure_ascii=False)


def create_project_structure(project_name, project_type, base_dir, custom_dirs=None,
                             create_readme=True, create_main_py=True, services=None):
    """
    Cria a estrutura de diretórios para um novo projeto
    Retorna o caminho completo do projeto criado
    """
    project_path = os.path.join(base_dir, project_name)

    # Verifica se o projeto já existe e adiciona sufixo se necessário
    if os.path.exists(project_path):
        counter = 1
        while os.path.exists(f"{project_path}_{counter}"):
            counter += 1
        project_path = f"{project_path}_{counter}"
        project_name = f"{project_name}_{counter}"

    # Cria os diretórios principais
    dirs = custom_dirs if custom_dirs else ['dados', 'documentos', 'relatorios', 'planilhas', 'scripts', 'outputs']
    os.makedirs(project_path)

    for dir_name in dirs:
        os.makedirs(os.path.join(project_path, dir_name), exist_ok=True)

    # Cria arquivos iniciais
    if create_readme:
        create_readme(project_path, project_name, project_type)
    if create_main_py:
        create_main_py(project_path)
    if services:
        create_services_config(project_path, services)

    return project_path


def create_readme(project_path, project_name, project_type):
    """Cria um arquivo README.md com informações do projeto"""
    content = f"""# {project_name}

**Tipo de Projeto**: {project_type}

**Data de Criação**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## Descrição

Este projeto foi criado automaticamente pelo Project Wizard PRO MAX.

## Estrutura de Diretórios

- `/dados/`: Armazena arquivos de dados brutos
- `/documentos/`: Documentação do projeto
- `/relatorios/`: Relatórios e análises
- `/planilhas/`: Arquivos Excel e planilhas
- `/scripts/`: Scripts Python, Shell, etc.
- `/outputs/`: Resultados e saídas do projeto

## Como Usar

1. Descreva o propósito do projeto
2. Adicione instruções de configuração
3. Documente os principais scripts e funções
"""

    with open(os.path.join(project_path, 'README.md'), 'w', encoding='utf-8') as f:
        f.write(content)


def create_main_py(project_path):
    """Cria um arquivo main.py básico"""
    content = '''"""Script principal do projeto"""

def main():
    """Função principal do projeto"""
    print("Hello World!")

if __name__ == "__main__":
    main()
'''
    with open(os.path.join(project_path, 'scripts', 'main.py'), 'w', encoding='utf-8') as f:
        f.write(content)


def create_services_config(project_path, services):
    """Cria um arquivo de configuração para os serviços do projeto"""
    config = {
        "services": services,
        "created_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }

    with open(os.path.join(project_path, 'project_services.json'), 'w', encoding='utf-8') as f:
        json.dump(config, f, indent=4)


def backup_project(project_path):
    """Cria um backup zip do projeto"""
    backup_dir = os.path.join(project_path, 'backups')
    os.makedirs(backup_dir, exist_ok=True)

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    backup_file = os.path.join(backup_dir, f'backup_{timestamp}.zip')

    with zipfile.ZipFile(backup_file, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(project_path):
            # Ignora o próprio diretório de backups para evitar recursão
            if os.path.basename(root) == 'backups':
                continue

            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, start=project_path)
                zipf.write(file_path, arcname)

    return backup_file


def add_to_history(project_name, project_path, project_type):
    """Adiciona o projeto ao histórico"""
    try:
        with open('data/historico.json', 'r+', encoding='utf-8') as f:
            try:
                history = json.load(f)
            except json.JSONDecodeError:
                history = []

            history.append({
                'name': project_name,
                'path': project_path,
                'type': project_type,
                'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            })

            f.seek(0)
            json.dump(history, f, indent=4, ensure_ascii=False)
            f.truncate()
    except Exception as e:
        print(f"Erro ao adicionar ao histórico: {str(e)}")


def get_recent_projects(limit=5):
    """Retorna os projetos mais recentes"""
    try:
        with open('data/historico.json', 'r', encoding='utf-8') as f:
            history = json.load(f)
            return history[-limit:][::-1]  # Retorna os mais recentes primeiro
    except (FileNotFoundError, json.JSONDecodeError):
        return []


def add_note(title, content, tags=None, note_type="geral"):
    """Adiciona uma nova anotação"""
    try:
        with open('data/notes.json', 'r+', encoding='utf-8') as f:
            try:
                notes = json.load(f)
            except json.JSONDecodeError:
                notes = []

            notes.append({
                'id': len(notes) + 1,
                'title': title,
                'content': content,
                'tags': tags or [],
                'type': note_type,
                'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'updated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            })

            f.seek(0)
            json.dump(notes, f, indent=4, ensure_ascii=False)
            f.truncate()
    except Exception as e:
        print(f"Erro ao adicionar anotação: {str(e)}")


def get_notes(tag=None, note_type=None):
    """Retorna as anotações, filtradas por tag e tipo se especificado"""
    try:
        with open('data/notes.json', 'r', encoding='utf-8') as f:
            notes = json.load(f)

            if tag:
                notes = [note for note in notes if tag in note.get('tags', [])]

            if note_type:
                notes = [note for note in notes if note.get('type') == note_type]

            return notes
    except (FileNotFoundError, json.JSONDecodeError):
        return []


def export_notes_to_docx(notes, output_path):
    """Exporta anotações para um arquivo DOCX"""
    doc = Document()

    # Configuração básica do documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Adiciona título do documento
    doc.add_heading('Anotações Exportadas', level=0)
    doc.add_paragraph(f"Data de exportação: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("\n")

    # Adiciona cada anotação
    for note in notes:
        doc.add_heading(note['title'], level=1)
        doc.add_paragraph(f"Tipo: {note['type']}")
        doc.add_paragraph(f"Tags: {', '.join(note.get('tags', []))}")
        doc.add_paragraph(f"Criado em: {note['created_at']}")
        doc.add_paragraph("\n")

        # Adiciona o conteúdo da anotação
        for paragraph in note['content'].split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)

        doc.add_paragraph("\n")
        doc.add_paragraph("-" * 50)
        doc.add_paragraph("\n")

    # Salva o documento
    doc.save(output_path)


def clean_temp_files(directory):
    """Limpa arquivos temporários e lixo do diretório especificado"""
    temp_extensions = ['.tmp', '.temp', '.bak', '.~', '.log']
    removed_files = []

    for root, dirs, files in os.walk(directory):
        for file in files:
            if any(file.lower().endswith(ext) for ext in temp_extensions):
                try:
                    file_path = os.path.join(root, file)
                    os.remove(file_path)
                    removed_files.append(file_path)
                except Exception as e:
                    print(f"Erro ao remover {file}: {str(e)}")

    return removed_files