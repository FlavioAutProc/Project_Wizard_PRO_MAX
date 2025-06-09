import os
import re
import json
import sqlite3
import tempfile
import pytesseract
import numpy as np
import pandas as pd
import streamlit as st
from PIL import Image
import pdfplumber
import fitz  # PyMuPDF
import cv2
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from docx import Document
from fpdf import FPDF
import datetime
import random
from collections import defaultdict
from io import BytesIO

# Configuração do Tesseract OCR
try:
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
except:
    st.warning("Configure manualmente o caminho do Tesseract nas Configurações do sistema")

# Configurações iniciais
nltk.download('punkt')
nltk.download('stopwords')

# Configuração do Streamlit
st.set_page_config(
    page_title="EstudaZilla Ultimate v2.0",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Classes do sistema
class PDFProcessor:
    def __init__(self):
        self.text_blocks = []
        self.current_chapter = ""
        self.current_theme = ""
        self.current_subtheme = ""

    def process_pdf(self, file_path):
        """Processa um arquivo PDF, extraindo texto e imagens"""
        try:
            # Primeira passagem com pdfplumber para texto
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        self._process_text_block(text, page_num + 1)

            # Segunda passagem com PyMuPDF para OCR em imagens
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                # Verifica se a página é principalmente imagem
                if len(self.text_blocks) == 0 or page_num + 1 > len(self.text_blocks) or not self.text_blocks[page_num][
                    'text'].strip():
                    text = self._ocr_image(img)
                    if text.strip():
                        self._process_text_block(text, page_num + 1)

            return self._structure_content()
        except Exception as e:
            st.error(f"Erro ao processar PDF: {str(e)}")
            return None

    def _process_text_block(self, text, page_num):
        """Processa um bloco de texto, identificando estrutura"""
        # Identifica capítulos, temas e subtemas
        chapter_match = re.search(r'(CAP[ÍI]TULO|CHAPTER)\s*(\d+)[.:]\s*(.+)', text, re.IGNORECASE)
        theme_match = re.search(r'(TEMA|TOPIC)\s*(\d+)[.:]\s*(.+)', text, re.IGNORECASE)
        subtheme_match = re.search(r'(SUB?TEMA|SUBTOPIC)\s*(\d+)[.:]\s*(.+)', text, re.IGNORECASE)

        if chapter_match:
            self.current_chapter = chapter_match.group(3).strip()
            self.current_theme = ""
            self.current_subtheme = ""
        elif theme_match:
            self.current_theme = theme_match.group(3).strip()
            self.current_subtheme = ""
        elif subtheme_match:
            self.current_subtheme = subtheme_match.group(3).strip()

        # Adiciona o bloco de texto
        self.text_blocks.append({
            'page': page_num,
            'chapter': self.current_chapter,
            'theme': self.current_theme,
            'subtheme': self.current_subtheme,
            'text': text
        })

    def _ocr_image(self, image):
        """Realiza OCR em uma imagem"""
        try:
            # Pré-processamento da imagem com OpenCV
            img_array = np.array(image)
            gray = cv2.cvtColor(img_array, cv2.COLOR_BGR2GRAY)
            thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

            # OCR com Tesseract
            custom_config = r'--oem 3 --psm 6'
            text = pytesseract.image_to_string(thresh, config=custom_config, lang='por')
            return text
        except Exception as e:
            st.warning(f"Erro no OCR: {str(e)}")
            return ""

    def _structure_content(self):
        """Estrutura o conteúdo extraído em um formato organizado"""
        structured = {
            'chapters': defaultdict(lambda: {
                'themes': defaultdict(lambda: {
                    'subthemes': defaultdict(list)
                })
            })
        }

        for block in self.text_blocks:
            chapter = block['chapter'] or "Sem Capítulo"
            theme = block['theme'] or "Sem Tema"
            subtheme = block['subtheme'] or "Sem Subtema"

            structured['chapters'][chapter]['themes'][theme]['subthemes'][subtheme].append({
                'page': block['page'],
                'text': block['text']
            })

        return structured


class ContentSummarizer:
    def __init__(self):
        self.stop_words = set(stopwords.words('portuguese'))

    def generate_summary(self, text, style='bullet', sentences_count=5):
        """Gera um resumo do texto no estilo especificado"""
        if style == 'bullet':
            return self._bullet_summary(text, sentences_count)
        elif style == 'flashcard':
            return self._flashcard_summary(text)
        elif style == 'dissertative':
            return self._dissertative_summary(text)
        elif style == 'mindmap':
            return self._mindmap_summary(text)
        else:
            return self._bullet_summary(text, sentences_count)

    def _bullet_summary(self, text, sentences_count):
        """Resumo em tópicos"""
        parser = PlaintextParser.from_string(text, Tokenizer('portuguese'))
        summarizer = LsaSummarizer()
        summary = summarizer(parser.document, sentences_count)
        return "\n• ".join([str(sentence) for sentence in summary])

    def _flashcard_summary(self, text):
        """Resumo em formato de flashcards (pergunta e resposta)"""
        sentences = sent_tokenize(text, language='portuguese')
        if len(sentences) < 2:
            return "Pergunta: Qual é o tópico principal?\nResposta: " + text[:200] + "..."

        question = self._create_question(sentences[0])
        answer = " ".join(sentences[1:3])[:300] + "..."
        return f"Pergunta: {question}\nResposta: {answer}"

    def _dissertative_summary(self, text):
        """Resumo dissertativo"""
        sentences = sent_tokenize(text, language='portuguese')
        if not sentences:
            return ""

        keywords = self._extract_keywords(text)
        intro = f"O texto aborda principalmente sobre {', '.join(keywords[:3])}. "
        development = " ".join(sentences[:3])[:500] + "..."
        conclusion = "Portanto, pode-se compreender que " + sentences[-1][:150] + "..."

        return intro + development + conclusion

    def _mindmap_summary(self, text):
        """Resumo em formato de mapa mental"""
        keywords = self._extract_keywords(text)
        main_topic = keywords[0] if keywords else "Tópico Principal"

        branches = []
        for i, kw in enumerate(keywords[1:4], 1):
            branches.append(f"    └── {kw}")

        return f"{main_topic}\n" + "\n".join(branches)

    def _extract_keywords(self, text, num_keywords=5):
        """Extrai palavras-chave do texto"""
        words = word_tokenize(text.lower(), language='portuguese')
        words = [word for word in words if word.isalnum() and word not in self.stop_words]
        freq = nltk.FreqDist(words)
        return [word for word, _ in freq.most_common(num_keywords)]

    def _create_question(self, sentence):
        """Cria uma pergunta a partir de uma afirmação"""
        sentence = sentence.replace('é', 'foi').replace('são', 'foram')
        return sentence[:100] + "?"


class QuizGenerator:
    def __init__(self):
        self.question_types = {
            'multiple_choice': self._generate_multiple_choice,
            'true_false': self._generate_true_false,
            'short_answer': self._generate_short_answer,
            'case_study': self._generate_case_study
        }

    def generate_quiz(self, content, num_questions=5, q_type='multiple_choice'):
        """Gera um quiz com base no conteúdo"""
        generator = self.question_types.get(q_type, self._generate_multiple_choice)
        questions = []

        for _ in range(num_questions):
            question = generator(content)
            if question:
                questions.append(question)

        return questions

    def _generate_multiple_choice(self, content):
        """Gera questão de múltipla escolha"""
        sentences = self._extract_sentences(content)
        if len(sentences) < 4:
            return None

        question = self._create_question(sentences[0])
        correct = sentences[1][:100]
        options = [
            correct,
            sentences[2][:100],
            sentences[3][:100],
            self._modify_sentence(correct)
        ]
        random.shuffle(options)

        return {
            'type': 'multiple_choice',
            'question': question,
            'options': options,
            'answer': correct,
            'explanation': "Esta informação pode ser encontrada no texto original."
        }

    def _generate_true_false(self, content):
        """Gera questão de verdadeiro ou falso"""
        sentences = self._extract_sentences(content)
        if not sentences:
            return None

        sentence = random.choice(sentences)
        is_true = random.choice([True, False])
        modified_sentence = self._modify_sentence(sentence) if not is_true else sentence

        return {
            'type': 'true_false',
            'statement': modified_sentence,
            'answer': is_true,
            'explanation': "Esta afirmação está de acordo com o texto original." if is_true else "Esta afirmação contradiz o texto original."
        }

    def _generate_short_answer(self, content):
        """Gera questão de resposta curta"""
        sentences = self._extract_sentences(content)
        if not sentences:
            return None

        question = self._create_question(sentences[0])
        answer = sentences[1][:150]

        return {
            'type': 'short_answer',
            'question': question,
            'answer': answer,
            'explanation': "A resposta pode ser encontrada no texto original."
        }

    def _generate_case_study(self, content):
        """Gera estudo de caso"""
        sentences = self._extract_sentences(content)
        if len(sentences) < 3:
            return None

        context = " ".join(sentences[:3])
        question = "Como você resolveria esta situação?"
        answer = "Uma possível solução seria " + sentences[-1][:150]

        return {
            'type': 'case_study',
            'context': context,
            'question': question,
            'answer': answer,
            'explanation': "Esta é uma das possíveis abordagens baseadas no conteúdo estudado."
        }

    def _extract_sentences(self, content):
        """Extrai frases do conteúdo"""
        if isinstance(content, str):
            return sent_tokenize(content, language='portuguese')
        elif isinstance(content, dict):
            all_sentences = []
            for chapter in content.get('chapters', {}).values():
                for theme in chapter.get('themes', {}).values():
                    for subtheme in theme.get('subthemes', {}).values():
                        for block in subtheme:
                            all_sentences.extend(sent_tokenize(block['text'], language='portuguese'))
            return all_sentences
        return []

    def _create_question(self, sentence):
        """Cria uma pergunta a partir de uma afirmação"""
        sentence = sentence.replace('é', 'foi').replace('são', 'foram')
        return "Qual das alternativas abaixo completa corretamente: " + sentence[:100] + "..."

    def _modify_sentence(self, sentence):
        """Modifica uma frase para criar alternativas incorretas"""
        words = word_tokenize(sentence, language='portuguese')
        if len(words) < 3:
            return sentence + " (incorreta)"

        # Substitui algumas palavras por antônimos ou termos incorretos
        replacements = {
            'grande': 'pequeno',
            'aumento': 'diminuição',
            'positivo': 'negativo',
            'melhor': 'pior',
            'certo': 'errado'
        }

        for i, word in enumerate(words):
            if word.lower() in replacements:
                words[i] = replacements[word.lower()]
                break

        return " ".join(words) + "..."


class DatabaseManager:
    def __init__(self, db_path='estudazilla.db'):
        self.conn = sqlite3.connect(db_path)
        self._create_tables()

    def _create_tables(self):
        """Cria as tabelas do banco de dados"""
        cursor = self.conn.cursor()

        # Tabela de documentos
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            file_path TEXT NOT NULL,
            upload_date TEXT NOT NULL,
            last_accessed TEXT,
            category TEXT,
            pages INTEGER
        )
        ''')

        # Tabela de conteúdo
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS content (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            document_id INTEGER,
            chapter TEXT,
            theme TEXT,
            subtheme TEXT,
            page INTEGER,
            text_content TEXT,
            is_important INTEGER DEFAULT 0,
            last_reviewed TEXT,
            FOREIGN KEY (document_id) REFERENCES documents (id)
        )
        ''')

        # Tabela de flashcards
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS flashcards (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            content_id INTEGER,
            question TEXT NOT NULL,
            answer TEXT NOT NULL,
            created_date TEXT NOT NULL,
            last_reviewed TEXT,
            difficulty INTEGER DEFAULT 1,
            FOREIGN KEY (content_id) REFERENCES content (id)
        )
        ''')

        # Tabela de questões
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            content_id INTEGER,
            question_type TEXT NOT NULL,
            question_text TEXT NOT NULL,
            options TEXT,
            correct_answer TEXT NOT NULL,
            explanation TEXT,
            difficulty INTEGER DEFAULT 1,
            created_date TEXT NOT NULL,
            FOREIGN KEY (content_id) REFERENCES content (id)
        )
        ''')

        self.conn.commit()

    def save_document(self, title, file_path, category=None):
        """Salva um documento no banco de dados"""
        cursor = self.conn.cursor()
        cursor.execute('''
        INSERT INTO documents (title, file_path, upload_date, category)
        VALUES (?, ?, datetime('now'), ?)
        ''', (title, file_path, category))
        self.conn.commit()
        return cursor.lastrowid

    def save_content(self, document_id, structured_content):
        """Salva o conteúdo estruturado no banco de dados"""
        cursor = self.conn.cursor()

        for chapter, chapter_data in structured_content.get('chapters', {}).items():
            for theme, theme_data in chapter_data.get('themes', {}).items():
                for subtheme, blocks in theme_data.get('subthemes', {}).items():
                    for block in blocks:
                        cursor.execute('''
                        INSERT INTO content (document_id, chapter, theme, subtheme, page, text_content)
                        VALUES (?, ?, ?, ?, ?, ?)
                        ''', (document_id, chapter, theme, subtheme, block['page'], block['text']))

        self.conn.commit()

    def get_documents(self):
        """Obtém todos os documentos"""
        cursor = self.conn.cursor()
        cursor.execute('SELECT id, title, upload_date, category FROM documents ORDER BY last_accessed DESC')
        return cursor.fetchall()

    def get_document_content(self, document_id):
        """Obtém o conteúdo de um documento"""
        cursor = self.conn.cursor()
        cursor.execute('''
        SELECT id, chapter, theme, subtheme, page, text_content 
        FROM content 
        WHERE document_id = ?
        ORDER BY page
        ''', (document_id,))
        return cursor.fetchall()

    def save_flashcard(self, content_id, question, answer):
        """Salva um flashcard"""
        cursor = self.conn.cursor()
        cursor.execute('''
        INSERT INTO flashcards (content_id, question, answer, created_date)
        VALUES (?, ?, ?, datetime('now'))
        ''', (content_id, question, answer))
        self.conn.commit()
        return cursor.lastrowid

    def get_flashcards(self):
        """Obtém todos os flashcards"""
        cursor = self.conn.cursor()
        cursor.execute('''
        SELECT f.id, f.question, f.answer, c.chapter, c.theme 
        FROM flashcards f
        JOIN content c ON f.content_id = c.id
        ORDER BY f.last_reviewed ASC
        ''')
        return cursor.fetchall()

    def save_question(self, content_id, question_data):
        """Salva uma questão no banco de dados"""
        cursor = self.conn.cursor()
        options = json.dumps(question_data.get('options', [])) if 'options' in question_data else None

        cursor.execute('''
        INSERT INTO questions (
            content_id, question_type, question_text, options, 
            correct_answer, explanation, difficulty, created_date
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, datetime('now'))
        ''', (
            content_id,
            question_data['type'],
            question_data['question'],
            options,
            question_data['answer'],
            question_data.get('explanation', ''),
            question_data.get('difficulty', 1)
        ))
        self.conn.commit()
        return cursor.lastrowid

    def get_questions(self, question_type=None):
        """Obtém questões do banco de dados"""
        cursor = self.conn.cursor()

        if question_type:
            cursor.execute('''
            SELECT q.id, q.question_text, q.options, q.correct_answer, q.explanation, c.chapter, c.theme
            FROM questions q
            JOIN content c ON q.content_id = c.id
            WHERE q.question_type = ?
            ''', (question_type,))
        else:
            cursor.execute('''
            SELECT q.id, q.question_text, q.options, q.correct_answer, q.explanation, c.chapter, c.theme
            FROM questions q
            JOIN content c ON q.content_id = c.id
            ''')

        questions = []
        for row in cursor.fetchall():
            options = json.loads(row[2]) if row[2] else None
            questions.append({
                'id': row[0],
                'question': row[1],
                'options': options,
                'answer': row[3],
                'explanation': row[4],
                'chapter': row[5],
                'theme': row[6]
            })

        return questions

    def close(self):
        """Fecha a conexão com o banco de dados"""
        self.conn.close()


class ExportManager:
    def __init__(self):
        pass

    def export_txt(self, content, filename):
        """Exporta conteúdo para TXT"""
        with open(filename, 'w', encoding='utf-8') as f:
            if isinstance(content, dict):
                for chapter, chapter_data in content.get('chapters', {}).items():
                    f.write(f"CAPÍTULO: {chapter}\n\n")
                    for theme, theme_data in chapter_data.get('themes', {}).items():
                        f.write(f"TEMA: {theme}\n")
                        for subtheme, blocks in theme_data.get('subthemes', {}).items():
                            f.write(f"Subtema: {subtheme}\n")
                            for block in blocks:
                                f.write(f"Página {block['page']}:\n{block['text']}\n\n")
            else:
                f.write(content)
        return filename

    def export_pdf(self, content, filename):
        """Exporta conteúdo para PDF"""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        if isinstance(content, dict):
            for chapter, chapter_data in content.get('chapters', {}).items():
                pdf.cell(200, 10, txt=f"CAPÍTULO: {chapter}", ln=True)
                pdf.ln(5)

                for theme, theme_data in chapter_data.get('themes', {}).items():
                    pdf.cell(200, 10, txt=f"TEMA: {theme}", ln=True)

                    for subtheme, blocks in theme_data.get('subthemes', {}).items():
                        pdf.cell(200, 10, txt=f"Subtema: {subtheme}", ln=True)

                        for block in blocks:
                            pdf.multi_cell(0, 10, txt=f"Página {block['page']}:\n{block['text']}")
                            pdf.ln(5)
        else:
            pdf.multi_cell(0, 10, txt=content)

        pdf.output(filename)
        return filename

    def export_docx(self, content, filename):
        """Exporta conteúdo para DOCX"""
        doc = Document()

        if isinstance(content, dict):
            for chapter, chapter_data in content.get('chapters', {}).items():
                doc.add_heading(f"CAPÍTULO: {chapter}", level=1)

                for theme, theme_data in chapter_data.get('themes', {}).items():
                    doc.add_heading(f"TEMA: {theme}", level=2)

                    for subtheme, blocks in theme_data.get('subthemes', {}).items():
                        doc.add_heading(f"Subtema: {subtheme}", level=3)

                        for block in blocks:
                            doc.add_paragraph(f"Página {block['page']}:")
                            doc.add_paragraph(block['text'])
                            doc.add_paragraph()
        else:
            doc.add_paragraph(content)

        doc.save(filename)
        return filename

    def export_quiz(self, questions, filename, format='pdf'):
        """Exporta um quiz para o formato especificado"""
        if format == 'pdf':
            return self._export_quiz_pdf(questions, filename)
        elif format == 'docx':
            return self._export_quiz_docx(questions, filename)
        else:
            return self._export_quiz_txt(questions, filename)

    def _export_quiz_pdf(self, questions, filename):
        """Exporta quiz para PDF"""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        pdf.cell(200, 10, txt="SIMULADO GERADO PELO ESTUDAZILLA", ln=True)
        pdf.ln(10)

        for i, question in enumerate(questions, 1):
            pdf.cell(200, 10, txt=f"{i}. {question['question']}", ln=True)

            if question.get('options'):
                for j, option in enumerate(question['options'], 1):
                    pdf.cell(200, 10, txt=f"   {chr(96 + j)}) {option}", ln=True)

            pdf.ln(5)

        # Adiciona gabarito em uma nova página
        pdf.add_page()
        pdf.cell(200, 10, txt="GABARITO", ln=True)
        pdf.ln(10)

        for i, question in enumerate(questions, 1):
            answer = question['answer']
            if question.get('options'):
                # Encontra o índice da resposta correta
                try:
                    idx = question['options'].index(answer)
                    answer = f"{chr(97 + idx)}) {answer}"
                except ValueError:
                    pass

            pdf.cell(200, 10, txt=f"{i}. {answer}", ln=True)
            if 'explanation' in question:
                pdf.multi_cell(0, 10, txt=f"   Explicação: {question['explanation']}")
            pdf.ln(5)

        pdf.output(filename)
        return filename

    def _export_quiz_docx(self, questions, filename):
        """Exporta quiz para DOCX"""
        doc = Document()
        doc.add_heading('SIMULADO GERADO PELO ESTUDAZILLA', level=1)

        for i, question in enumerate(questions, 1):
            doc.add_paragraph(f"{i}. {question['question']}")

            if question.get('options'):
                for j, option in enumerate(question['options'], 1):
                    doc.add_paragraph(f"   {chr(96 + j)}) {option}", style='ListBullet')

            doc.add_paragraph()

        # Adiciona gabarito
        doc.add_page_break()
        doc.add_heading('GABARITO', level=1)

        for i, question in enumerate(questions, 1):
            answer = question['answer']
            if question.get('options'):
                # Encontra o índice da resposta correta
                try:
                    idx = question['options'].index(answer)
                    answer = f"{chr(97 + idx)}) {answer}"
                except ValueError:
                    pass

            doc.add_paragraph(f"{i}. {answer}")
            if 'explanation' in question:
                doc.add_paragraph(f"   Explicação: {question['explanation']}")
            doc.add_paragraph()

        doc.save(filename)
        return filename

    def _export_quiz_txt(self, questions, filename):
        """Exporta quiz para TXT"""
        with open(filename, 'w', encoding='utf-8') as f:
            f.write("SIMULADO GERADO PELO ESTUDAZILLA\n\n")

            for i, question in enumerate(questions, 1):
                f.write(f"{i}. {question['question']}\n")

                if question.get('options'):
                    for j, option in enumerate(question['options'], 1):
                        f.write(f"   {chr(96 + j)}) {option}\n")

                f.write("\n")

            # Adiciona gabarito
            f.write("\n\nGABARITO\n\n")

            for i, question in enumerate(questions, 1):
                answer = question['answer']
                if question.get('options'):
                    # Encontra o índice da resposta correta
                    try:
                        idx = question['options'].index(answer)
                        answer = f"{chr(97 + idx)}) {answer}"
                    except ValueError:
                        pass

                f.write(f"{i}. {answer}\n")
                if 'explanation' in question:
                    f.write(f"   Explicação: {question['explanation']}\n")
                f.write("\n")

        return filename


# Interface do Streamlit
class EstudaZillaUI:
    def __init__(self):
        self.db = DatabaseManager()
        self.processor = PDFProcessor()
        self.summarizer = ContentSummarizer()
        self.quiz_generator = QuizGenerator()
        self.exporter = ExportManager()

        # Configuração do estado da sessão
        if 'current_document' not in st.session_state:
            st.session_state.current_document = None
        if 'current_content' not in st.session_state:
            st.session_state.current_content = None
        if 'flashcards' not in st.session_state:
            st.session_state.flashcards = []
        if 'questions' not in st.session_state:
            st.session_state.questions = []
        if 'pomodoro_active' not in st.session_state:
            st.session_state.pomodoro_active = False
        if 'pomodoro_start' not in st.session_state:
            st.session_state.pomodoro_start = None
        if 'pomodoro_duration' not in st.session_state:
            st.session_state.pomodoro_duration = 25

    def show_sidebar(self):
        """Mostra a barra lateral com navegação e informações"""
        with st.sidebar:
            st.image("https://via.placeholder.com/150x50?text=EstudaZilla", use_column_width=True)

            menu_options = {
                "📁 Documentos": self.show_documents_tab,
                "📄 Resumos": self.show_summaries_tab,
                "🎓 Simulados": self.show_quizzes_tab,
                "🔁 Flashcards": self.show_flashcards_tab,
                "⚙️ Configurações": self.show_settings_tab
            }

            selected = st.radio(
                "Menu",
                list(menu_options.keys()),
                index=0
            )

            # Mostra o Pomodoro timer
            self._show_pomodoro_timer()

            # Mostra estatísticas
            self._show_stats()

        # Mostra a aba selecionada
        menu_options[selected]()

    def _show_pomodoro_timer(self):
        """Mostra o timer Pomodoro na barra lateral"""
        with st.sidebar.expander("⏱️ Pomodoro Timer"):
            col1, col2 = st.columns(2)

            with col1:
                duration = st.select_slider(
                    "Duração (min)",
                    options=[15, 20, 25, 30, 45, 60],
                    value=st.session_state.pomodoro_duration
                )
                st.session_state.pomodoro_duration = duration

            with col2:
                if st.button("Iniciar" if not st.session_state.pomodoro_active else "Parar"):
                    st.session_state.pomodoro_active = not st.session_state.pomodoro_active
                    if st.session_state.pomodoro_active:
                        st.session_state.pomodoro_start = datetime.datetime.now()
                    else:
                        st.session_state.pomodoro_start = None

            if st.session_state.pomodoro_active and st.session_state.pomodoro_start:
                elapsed = (datetime.datetime.now() - st.session_state.pomodoro_start).total_seconds()
                remaining = max(0, st.session_state.pomodoro_duration * 60 - elapsed)

                minutes, seconds = divmod(int(remaining), 60)
                st.metric("Tempo Restante", f"{minutes:02d}:{seconds:02d}")

                # Barra de progresso
                progress = min(1.0, elapsed / (st.session_state.pomodoro_duration * 60))
                st.progress(progress)

                # Verifica se o tempo acabou
                if remaining <= 0:
                    st.session_state.pomodoro_active = False
                    st.session_state.pomodoro_start = None
                    st.warning("Tempo esgotado! Hora de uma pausa.")
                    st.balloons()

    def _show_stats(self):
        """Mostra estatísticas na barra lateral"""
        with st.sidebar.expander("📊 Estatísticas"):
            cursor = self.db.conn.cursor()

            # Documentos
            cursor.execute('SELECT COUNT(*) FROM documents')
            doc_count = cursor.fetchone()[0]

            # Flashcards
            cursor.execute('SELECT COUNT(*) FROM flashcards')
            flashcard_count = cursor.fetchone()[0]

            # Questões
            cursor.execute('SELECT COUNT(*) FROM questions')
            question_count = cursor.fetchone()[0]

            st.metric("Documentos", doc_count)
            st.metric("Flashcards", flashcard_count)
            st.metric("Questões", question_count)

            # Sugestão de revisão
            cursor.execute('''
            SELECT c.chapter, c.theme, COUNT(*) as count
            FROM content c
            LEFT JOIN flashcards f ON c.id = f.content_id
            WHERE f.id IS NULL
            GROUP BY c.chapter, c.theme
            ORDER BY count DESC
            LIMIT 1
            ''')
            suggestion = cursor.fetchone()

            if suggestion:
                st.write("**Sugestão de Revisão**")
                st.write(f"{suggestion[0]} - {suggestion[1]}")
                st.write(f"{suggestion[2]} trechos sem flashcards")

    def show_documents_tab(self):
        """Mostra a aba de documentos"""
        st.title("📁 Documentos")

        # Upload de arquivos
        with st.expander("⬆️ Upload de Documentos", expanded=True):
            uploaded_files = st.file_uploader(
                "Carregue seus PDFs",
                type=['pdf'],
                accept_multiple_files=True
            )

            if uploaded_files:
                progress_bar = st.progress(0)
                status_text = st.empty()

                for i, uploaded_file in enumerate(uploaded_files):
                    try:
                        # Salva o arquivo temporariamente
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                            tmp_file.write(uploaded_file.getvalue())
                            tmp_path = tmp_file.name

                        # Processa o PDF
                        status_text.text(f"Processando {uploaded_file.name}...")
                        structured_content = self.processor.process_pdf(tmp_path)

                        if structured_content:
                            # Salva no banco de dados
                            doc_id = self.db.save_document(uploaded_file.name, tmp_path)
                            self.db.save_content(doc_id, structured_content)

                            # Atualiza a interface
                            st.session_state.current_document = doc_id
                            st.session_state.current_content = structured_content

                            st.success(f"Documento {uploaded_file.name} processado com sucesso!")
                        else:
                            st.error(f"Falha ao processar {uploaded_file.name}")

                        # Remove o arquivo temporário
                        os.unlink(tmp_path)

                        # Atualiza a barra de progresso
                        progress_bar.progress((i + 1) / len(uploaded_files))

                    except Exception as e:
                        st.error(f"Erro ao processar {uploaded_file.name}: {str(e)}")

                status_text.text("Processamento concluído!")

        # Lista de documentos
        with st.expander("📂 Documentos Armazenados"):
            documents = self.db.get_documents()

            if documents:
                cols = st.columns([3, 2, 2, 1])
                cols[0].write("**Título**")
                cols[1].write("**Data de Upload**")
                cols[2].write("**Categoria**")
                cols[3].write("**Ações**")

                for doc in documents:
                    doc_id, title, upload_date, category = doc
                    cols = st.columns([3, 2, 2, 1])

                    cols[0].write(title)
                    cols[1].write(upload_date)
                    cols[2].write(category or "Sem categoria")

                    if cols[3].button("Abrir", key=f"open_{doc_id}"):
                        st.session_state.current_document = doc_id
                        content = self.db.get_document_content(doc_id)

                        # Converte para o formato estruturado
                        structured = {'chapters': defaultdict(
                            lambda: {'themes': defaultdict(lambda: {'subthemes': defaultdict(list)})})}

                        for item in content:
                            _, chapter, theme, subtheme, page, text = item
                            structured['chapters'][chapter]['themes'][theme]['subthemes'][subtheme].append({
                                'page': page,
                                'text': text
                            })

                        st.session_state.current_content = structured
                        st.rerun()
            else:
                st.info("Nenhum documento carregado ainda.")

        # Visualizador de conteúdo
        if st.session_state.current_content:
            self._show_document_content(st.session_state.current_content)

    def _show_document_content(self, content):
        """Mostra o conteúdo de um documento"""
        st.write("📄 Visualizador de Conteúdo")

        chapters = list(content['chapters'].keys())

        if chapters:
            selected_chapter = st.selectbox(
                "Capítulo",
                chapters,
                index=0
            )

            chapter_data = content['chapters'][selected_chapter]
            themes = list(chapter_data['themes'].keys())

            if themes:
                selected_theme = st.selectbox(
                    "Tema",
                    themes,
                    index=0
                )

                theme_data = chapter_data['themes'][selected_theme]
                subthemes = list(theme_data['subthemes'].keys())

                if subthemes:
                    selected_subtheme = st.selectbox(
                        "Subtema",
                        subthemes,
                        index=0
                    )

                    blocks = theme_data['subthemes'][selected_subtheme]

                    for block in blocks:
                        st.subheader(f"Página {block['page']}")
                        st.text_area(
                            "Conteúdo",
                            value=block['text'],
                            height=200,
                            key=f"content_{block['page']}",
                            label_visibility="collapsed"
                        )

                        # Opções para o bloco de texto
                        col1, col2, col3 = st.columns([2, 2, 1])

                        with col1:
                            if st.button("Gerar Resumo", key=f"summary_{block['page']}"):
                                st.session_state.current_block = block['text']
                                st.session_state.show_summary = True

                        with col2:
                            if st.button("Gerar Questões", key=f"quiz_{block['page']}"):
                                st.session_state.current_block = block['text']
                                st.session_state.show_quiz = True

                        with col3:
                            if st.button("Flashcard", key=f"flash_{block['page']}"):
                                summary = self.summarizer.generate_summary(block['text'], style='flashcard')
                                question, answer = summary.split("\n")[0], "\n".join(summary.split("\n")[1:])
                                self.db.save_flashcard(st.session_state.current_document, question, answer)
                                st.success("Flashcard criado com sucesso!")

            # Mostra resumo se solicitado
            if hasattr(st.session_state, 'show_summary') and st.session_state.show_summary:
                self._show_summary_options(st.session_state.current_block)

            # Mostra questões se solicitado
            if hasattr(st.session_state, 'show_quiz') and st.session_state.show_quiz:
                self._show_quiz_options(st.session_state.current_block)

    def _show_summary_options(self, text):
        """Mostra opções para geração de resumo"""
        with st.expander("📝 Gerar Resumo", expanded=True):
            summary_type = st.radio(
                "Tipo de Resumo",
                ['Bullet Points ✅', 'Flashcards 🔁', 'Dissertativo 📝', 'Mapa Mental 🧠'],
                index=0
            )

            style_map = {
                'Bullet Points ✅': 'bullet',
                'Flashcards 🔁': 'flashcard',
                'Dissertativo 📝': 'dissertative',
                'Mapa Mental 🧠': 'mindmap'
            }

            if st.button("Gerar"):
                summary = self.summarizer.generate_summary(text, style=style_map[summary_type])

                st.subheader("Resumo Gerado")
                st.text_area("Resumo", value=summary, height=300)

                # Opções de exportação
                export_format = st.radio(
                    "Exportar como",
                    ['TXT', 'PDF', 'DOCX'],
                    horizontal=True
                )

                if st.button("Exportar"):
                    with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{export_format.lower()}') as tmp_file:
                        filename = tmp_file.name

                    if export_format == 'TXT':
                        self.exporter.export_txt(summary, filename)
                    elif export_format == 'PDF':
                        self.exporter.export_pdf(summary, filename)
                    elif export_format == 'DOCX':
                        self.exporter.export_docx(summary, filename)

                    with open(filename, 'rb') as f:
                        st.download_button(
                            "Baixar Arquivo",
                            f.read(),
                            file_name=f"resumo_estudazilla.{export_format.lower()}",
                            mime=f"application/{export_format.lower()}"
                        )

                    os.unlink(filename)

            if st.button("Fechar"):
                del st.session_state.show_summary
                st.rerun()

    def _show_quiz_options(self, text):
        """Mostra opções para geração de quiz"""
        with st.expander("🧪 Gerar Questões", expanded=True):
            question_type = st.radio(
                "Tipo de Questão",
                ['Múltipla Escolha', 'Verdadeiro/Falso', 'Resposta Curta', 'Estudo de Caso'],
                index=0
            )

            num_questions = st.slider("Número de Questões", 1, 10, 3)

            type_map = {
                'Múltipla Escolha': 'multiple_choice',
                'Verdadeiro/Falso': 'true_false',
                'Resposta Curta': 'short_answer',
                'Estudo de Caso': 'case_study'
            }

            if st.button("Gerar"):
                questions = self.quiz_generator.generate_quiz(
                    text,
                    num_questions=num_questions,
                    q_type=type_map[question_type]
                )

                st.session_state.generated_questions = questions
                st.session_state.show_quiz_results = True

            if hasattr(st.session_state, 'show_quiz_results') and st.session_state.show_quiz_results:
                self._show_quiz_results(st.session_state.generated_questions)

            if st.button("Fechar"):
                if hasattr(st.session_state, 'show_quiz_results'):
                    del st.session_state.show_quiz_results
                del st.session_state.show_quiz
                st.rerun()

    def _show_quiz_results(self, questions):
        """Mostra os resultados do quiz gerado"""
        st.subheader("Questões Geradas")

        for i, question in enumerate(questions, 1):
            with st.expander(f"Questão {i}"):
                st.write(f"**{question['question']}**")

                if question.get('options'):
                    st.write("**Opções:**")
                    for j, option in enumerate(question['options'], 1):
                        st.write(f"{chr(96 + j)}) {option}")

                with st.expander("Ver Resposta"):
                    st.write(f"**Resposta:** {question['answer']}")
                    if 'explanation' in question:
                        st.write(f"**Explicação:** {question['explanation']}")

                if st.button(f"Salvar Questão {i}"):
                    self.db.save_question(st.session_state.current_document, question)
                    st.success("Questão salva com sucesso!")

        # Opções de exportação
        st.subheader("Exportar Simulado")
        export_format = st.radio(
            "Formato",
            ['PDF', 'DOCX', 'TXT'],
            horizontal=True,
            key='quiz_export_format'
        )

        if st.button("Exportar Todas as Questões"):
            with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{export_format.lower()}') as tmp_file:
                filename = tmp_file.name

            self.exporter.export_quiz(questions, filename, format=export_format.lower())

            with open(filename, 'rb') as f:
                st.download_button(
                    "Baixar Simulado",
                    f.read(),
                    file_name=f"simulado_estudazilla.{export_format.lower()}",
                    mime=f"application/{export_format.lower()}"
                )

            os.unlink(filename)

    def show_summaries_tab(self):
        """Mostra a aba de resumos"""
        st.title("📄 Resumos")
        st.info("Esta funcionalidade está em desenvolvimento. Use a aba de Documentos para gerar resumos.")

    def show_quizzes_tab(self):
        """Mostra a aba de simulados"""
        st.title("🎓 Simulados")

        # Banco de questões
        with st.expander("📚 Banco de Questões"):
            question_types = ['Todas', 'Múltipla Escolha', 'Verdadeiro/Falso', 'Resposta Curta', 'Estudo de Caso']
            selected_type = st.selectbox("Filtrar por tipo", question_types, index=0)

            type_map = {
                'Todas': None,
                'Múltipla Escolha': 'multiple_choice',
                'Verdadeiro/Falso': 'true_false',
                'Resposta Curta': 'short_answer',
                'Estudo de Caso': 'case_study'
            }

            questions = self.db.get_questions(type_map[selected_type])

            if questions:
                for question in questions:
                    with st.expander(f"Questão {question['id']}"):
                        st.write(f"**{question['question']}**")

                        if question.get('options'):
                            st.write("**Opções:**")
                            for j, option in enumerate(question['options'], 1):
                                st.write(f"{chr(96 + j)}) {option}")

                        st.write(f"**Capítulo:** {question['chapter']}")
                        st.write(f"**Tema:** {question['theme']}")

                        with st.expander("Ver Resposta"):
                            st.write(f"**Resposta:** {question['answer']}")
                            if 'explanation' in question:
                                st.write(f"**Explicação:** {question['explanation']}")
            else:
                st.info("Nenhuma questão encontrada. Gere questões na aba de Documentos.")

        # Criar simulado personalizado
        with st.expander("✏️ Criar Simulado Personalizado"):
            st.warning("Esta funcionalidade está em desenvolvimento.")

    def show_flashcards_tab(self):
        """Mostra a aba de flashcards"""
        st.title("🔁 Flashcards")

        # Obter flashcards do banco de dados
        flashcards = self.db.get_flashcards()
        st.session_state.flashcards = flashcards

        if flashcards:
            # Controles de navegação
            if 'current_flashcard' not in st.session_state:
                st.session_state.current_flashcard = 0
                st.session_state.show_answer = False

            total = len(flashcards)
            current = st.session_state.current_flashcard + 1

            col1, col2, col3 = st.columns([1, 2, 1])
            with col1:
                if st.button("⏮️ Anterior") and st.session_state.current_flashcard > 0:
                    st.session_state.current_flashcard -= 1
                    st.session_state.show_answer = False
                    st.rerun()

            with col2:
                st.write(f"Flashcard {current} de {total}")
                st.write(
                    f"**Tópico:** {flashcards[st.session_state.current_flashcard][3]} - {flashcards[st.session_state.current_flashcard][4]}")

            with col3:
                if st.button("⏭️ Próximo") and st.session_state.current_flashcard < total - 1:
                    st.session_state.current_flashcard += 1
                    st.session_state.show_answer = False
                    st.rerun()

            # Mostra o flashcard atual
            card = st.container(border=True)
            card.write(f"**Pergunta:** {flashcards[st.session_state.current_flashcard][1]}")

            if st.session_state.show_answer:
                card.write(f"**Resposta:** {flashcards[st.session_state.current_flashcard][2]}")

                # Avaliação da dificuldade
                col1, col2, col3 = st.columns(3)
                with col1:
                    if st.button("Fácil 😊"):
                        self._update_flashcard_difficulty(flashcards[st.session_state.current_flashcard][0], 1)
                with col2:
                    if st.button("Médio 😐"):
                        self._update_flashcard_difficulty(flashcards[st.session_state.current_flashcard][0], 2)
                with col3:
                    if st.button("Difícil 😓"):
                        self._update_flashcard_difficulty(flashcards[st.session_state.current_flashcard][0], 3)
            else:
                if st.button("Mostrar Resposta"):
                    st.session_state.show_answer = True
                    st.rerun()

            # Exportar flashcards
            with st.expander("📤 Exportar Flashcards"):
                export_format = st.radio(
                    "Formato",
                    ['PDF', 'DOCX', 'TXT'],
                    horizontal=True
                )

                if st.button("Exportar Todos"):
                    with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{export_format.lower()}') as tmp_file:
                        filename = tmp_file.name

                    # Cria conteúdo para exportação
                    content = []
                    for card in flashcards:
                        content.append(f"Pergunta: {card[1]}\nResposta: {card[2]}\n\n")

                    if export_format == 'PDF':
                        self.exporter.export_pdf("".join(content), filename)
                    elif export_format == 'DOCX':
                        self.exporter.export_docx("".join(content), filename)
                    elif export_format == 'TXT':
                        self.exporter.export_txt("".join(content), filename)

                    with open(filename, 'rb') as f:
                        st.download_button(
                            "Baixar Flashcards",
                            f.read(),
                            file_name=f"flashcards_estudazilla.{export_format.lower()}",
                            mime=f"application/{export_format.lower()}"
                        )

                    os.unlink(filename)
        else:
            st.info("Nenhum flashcard encontrado. Crie flashcards na aba de Documentos.")

    def _update_flashcard_difficulty(self, flashcard_id, difficulty):
        """Atualiza a dificuldade de um flashcard"""
        cursor = self.db.conn.cursor()
        cursor.execute('''
        UPDATE flashcards 
        SET difficulty = ?, last_reviewed = datetime('now')
        WHERE id = ?
        ''', (difficulty, flashcard_id))
        self.db.conn.commit()
        st.success("Obrigado pelo feedback! Flashcard atualizado.")
        st.rerun()

    def show_settings_tab(self):
        """Mostra a aba de configurações"""
        st.title("⚙️ Configurações")

        with st.expander("🔧 Configurações do Sistema"):
            st.write("**Configurações de OCR**")
            tesseract_path = st.text_input("Caminho do Tesseract (opcional)", pytesseract.pytesseract.tesseract_cmd)

            if st.button("Salvar Configurações"):
                if tesseract_path and os.path.exists(tesseract_path):
                    pytesseract.pytesseract.tesseract_cmd = tesseract_path
                    st.success("Configurações salvas com sucesso!")
                else:
                    st.error("Caminho do Tesseract inválido. Usando padrão do sistema.")

        with st.expander("🗄️ Gerenciamento de Dados"):
            st.write("**Backup e Restauração**")

            col1, col2 = st.columns(2)
            with col1:
                if st.button("Exportar Banco de Dados"):
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.db') as tmp_file:
                        filename = tmp_file.name

                    # Cria uma cópia do banco de dados
                    self.db.conn.close()
                    import shutil
                    shutil.copy2('estudazilla.db', filename)
                    self.db = DatabaseManager()

                    with open(filename, 'rb') as f:
                        st.download_button(
                            "Baixar Backup",
                            f.read(),
                            file_name="estudazilla_backup.db",
                            mime="application/x-sqlite3"
                        )

                    os.unlink(filename)

            with col2:
                uploaded_db = st.file_uploader("Restaurar Backup", type=['db'])
                if uploaded_db and st.button("Restaurar"):
                    try:
                        self.db.conn.close()
                        with open('estudazilla.db', 'wb') as f:
                            f.write(uploaded_db.getvalue())

                        self.db = DatabaseManager()
                        st.success("Banco de dados restaurado com sucesso!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erro ao restaurar backup: {str(e)}")
                        self.db = DatabaseManager()

            if st.button("🔄 Redefinir Banco de Dados", type="primary"):
                if st.checkbox("Confirmar exclusão de TODOS os dados"):
                    try:
                        self.db.conn.close()
                        os.remove('estudazilla.db')
                        self.db = DatabaseManager()
                        st.session_state.current_document = None
                        st.session_state.current_content = None
                        st.success("Banco de dados redefinido com sucesso!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erro ao redefinir banco de dados: {str(e)}")
                        self.db = DatabaseManager()


# Função principal
def main():
    # Verifica se o Tesseract está instalado
    try:
        pytesseract.get_tesseract_version()
    except:
        st.warning("O Tesseract OCR não foi encontrado. A funcionalidade de OCR pode não funcionar corretamente.")

    # Inicializa a interface
    ui = EstudaZillaUI()
    ui.show_sidebar()


if __name__ == "__main__":
    main()